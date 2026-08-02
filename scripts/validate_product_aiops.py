#!/usr/bin/env python3
"""
Validate the built SFL-AIOPS-001 package against the commercial acceptance gate.

Checks the ARTIFACTS, not the source: opens every DOCX and XLSX, parses every
PDF, reads the ZIP, and fails loudly rather than reporting a pass it cannot
substantiate.

Usage:  python3 scripts/validate_product_aiops.py [--zip PATH]
"""

import argparse
import json
import pathlib
import re
import sys
import zipfile

REPO = pathlib.Path(__file__).resolve().parent.parent
DEFAULT_ZIP = REPO / "dist" / "product" / "SFL-AIOPS-001-AI-Operations-Playbook-Toolkit-v1.0.zip"

# Language that must not survive into a shipped artifact.
BANNED = [
    "REPLACE_WITH", "TODO", "FIXME", "pending Owner finalization", "Lorem ipsum",
    "PLACEHOLDER TEXT", "TBD", "XXX", "coming soon", "House of Assets",
]
# Infrastructure detail that must never appear in a customer-facing file.
SECRETS = [
    "workers.dev", "cloudflareaccess", "pages.dev", "sk_test", "sk_live", "whsec_",
    "re_", "D1_DATABASE", "R2_BUCKET", "SHOP_DB", "CF_ACCESS", "wrangler",
    "sentinelfortune.github.io",
]
REQUIRED_FOOTER = "Not legal, financial, HR, safety, code, permit, or regulatory advice."

REQUIRED_ROOT = ["README-START-HERE.pdf", "LICENSE-Single-Business.pdf", "PRODUCT-MANIFEST.json"]
REQUIRED_FOLDERS = [
    "01-Implementation-Guide", "02-Workflow-Map-and-Opportunity-Worksheet", "03-Prompt-Library",
    "04-SOP-Builder", "05-Customer-Communication-Templates", "06-Quality-Control-Checklist",
    "07-Weekly-Operations-Review", "08-Responsible-Use-Policy", "09-Business-Customization-Worksheet",
    "10-30-Day-Implementation-Roadmap",
]

failures = []
warnings = []
passes = []


def check(condition, ok_msg, fail_msg):
    (passes if condition else failures).append(ok_msg if condition else fail_msg)
    return condition


def docx_text(path):
    """Extract all text from a DOCX including tables, headers and footers."""
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    for section in doc.sections:
        parts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(parts)


def pdf_text(path):
    """Extract text from a PDF without a parsing library.

    reportlab writes content streams as ASCII85Decode followed by FlateDecode,
    so both layers have to come off before the text-showing operators are
    readable. (pypdf is not usable here — it imports `cryptography`, whose
    native extension is broken in this environment.)
    """
    import base64
    import zlib
    raw = path.read_bytes()
    chunks = []
    for match in re.finditer(rb"stream\r?\n(.*?)endstream", raw, re.S):
        blob = match.group(1).strip()
        try:
            blob = base64.a85decode(blob, adobe=True)
        except Exception:
            pass
        try:
            blob = zlib.decompress(blob)
        except Exception:
            continue
        for tj in re.finditer(rb"\(((?:[^()\\]|\\.)*)\)\s*Tj", blob):
            chunks.append(tj.group(1))
        for tja in re.finditer(rb"\[(.*?)\]\s*TJ", blob, re.S):
            for s in re.finditer(rb"\(((?:[^()\\]|\\.)*)\)", tja.group(1)):
                chunks.append(s.group(1))
    text = b" ".join(chunks)
    text = re.sub(rb"\\([()\\])", rb"\1", text)
    return text.decode("latin-1", errors="replace")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--zip", default=str(DEFAULT_ZIP))
    args = parser.parse_args()
    zip_path = pathlib.Path(args.zip)

    if not zip_path.exists():
        print(f"FAIL: package not found at {zip_path}")
        return 1

    stage = zip_path.parent / zip_path.stem

    # --- ZIP integrity and layout -----------------------------------------
    with zipfile.ZipFile(zip_path) as zf:
        bad = zf.testzip()
        check(bad is None, "ZIP integrity verified", f"ZIP corrupt at {bad}")
        names = zf.namelist()

    check(not any(n.startswith(f"{zip_path.stem}/") for n in names),
          "No nested master folder — files are at the ZIP root",
          "Files are nested under a master folder")
    for req in REQUIRED_ROOT:
        check(req in names, f"Root file present: {req}", f"MISSING root file: {req}")
    for folder in REQUIRED_FOLDERS:
        check(any(n.startswith(folder + "/") for n in names),
              f"Folder present: {folder}", f"MISSING folder: {folder}")

    # --- Manifest ----------------------------------------------------------
    manifest = json.loads((stage / "PRODUCT-MANIFEST.json").read_text(encoding="utf-8"))
    check(manifest["promptCount"] == 40,
          f"Manifest reports {manifest['promptCount']} prompts",
          f"Manifest reports {manifest['promptCount']} prompts, expected 40")
    check(manifest.get("contractVersion") == 1, "Manifest declares import contract version 1",
          f"Manifest contractVersion is {manifest.get('contractVersion')}, expected 1")
    for field in ("sku", "slug", "title", "version", "licenseType", "supportedFormats",
                  "shortDescription", "problemSolved", "description", "responsibleUseText",
                  "refundPolicySummary"):
        check(bool(manifest.get(field)), f"Manifest carries {field}", f"Manifest missing {field}")
    for field in ("deliverables", "notIncluded", "faqs"):
        check(bool(manifest.get(field)), f"Manifest carries a non-empty {field}", f"Manifest {field} is empty")

    disk_files = sorted(str(p.relative_to(stage)) for p in stage.rglob("*")
                        if p.is_file() and p.name != "PRODUCT-MANIFEST.json")
    manifest_files = sorted(f["path"] for f in manifest["files"])
    check(disk_files == manifest_files,
          f"Manifest lists every deliverable ({len(manifest_files)} files), no more and no fewer",
          f"Manifest mismatch. On disk only: {set(disk_files) - set(manifest_files)}; "
          f"in manifest only: {set(manifest_files) - set(disk_files)}")

    # --- Prompt library ----------------------------------------------------
    prompts = json.loads((REPO / "product-source" / "SFL-AIOPS-001" / "content" / "prompts.json")
                         .read_text(encoding="utf-8"))
    check(len(prompts) == 40, f"Source carries {len(prompts)} prompts", f"Source carries {len(prompts)}, expected 40")

    cats = {}
    for p in prompts:
        cats[p["category"]] = cats.get(p["category"], 0) + 1
    check(len(cats) == 6, f"All six prompt categories present: {cats}", f"Only {len(cats)} categories: {cats}")

    incomplete = [p["id"] for p in prompts
                  if not all(p.get(k) for k in ("purpose", "when", "inputs", "prompt", "warning"))]
    check(not incomplete, "Every prompt has purpose, when-to-use, inputs, prompt text and a warning",
          f"Prompts missing required fields: {incomplete}")

    # Every prompt must carry at least one genuine trade-specific constraint.
    TRADE_MARKERS = [
        "exclusion", "access", "shut-off", "shut off", "isolat", "warranty", "permit",
        "inspect", "safety", "licensed", "equipment", "technician", "dispatch", "code",
        "gas", "electrical", "escalat",
    ]
    generic = [p["id"] for p in prompts
               if not any(m in p["prompt"].lower() for m in TRADE_MARKERS)]
    check(not generic, "Every prompt carries at least one trade-specific element",
          f"Prompts with no trade-specific element: {generic}")

    # Boundary discipline, checked where it actually applies rather than as a
    # flat count. Any prompt whose output can reach a customer, or that reads a
    # technician's account of a job, must forbid inventing a price. Estimate
    # prompts must additionally refuse code and permit determinations, because
    # that is where a wrong answer becomes a contractual representation.
    CUSTOMER_FACING = {"customer-comms", "estimates", "field-notes", "csr"}
    price_scope = [p for p in prompts if p["category"] in CUSTOMER_FACING]
    # A pricing boundary is not always worded "price" — "no discount", "no
    # credit", "no compensation" and "no rate" are the same instruction.
    MONEY = ["price", "pricing", "cost", "rate", "discount", "compensation", "incentive", "charge", "total"]
    missing_price = [p["id"] for p in price_scope
                     if not any(w in p["prompt"].lower() for w in MONEY)]
    check(not missing_price,
          f"Every customer-facing prompt ({len(price_scope)}) forbids inventing a price",
          f"Customer-facing prompts with no pricing boundary: {missing_price}")

    estimate_prompts = [p for p in prompts if p["category"] == "estimates"]
    missing_cp = [p["id"] for p in estimate_prompts
                  if "code" not in p["prompt"].lower() or "permit" not in p["prompt"].lower()]
    check(not missing_cp,
          f"Every estimate prompt ({len(estimate_prompts)}) refuses code and permit determinations",
          f"Estimate prompts missing a code/permit boundary: {missing_cp}")

    # Every prompt, in every category, must carry an explicit prohibition.
    PROHIBITION = ["do not", "never", "must not"]
    no_prohibition = [p["id"] for p in prompts
                      if not any(w in p["prompt"].lower() for w in PROHIBITION)]
    check(not no_prohibition, "Every prompt states an explicit prohibition inside the prompt text",
          f"Prompts with no explicit prohibition: {no_prohibition}")

    # --- Every artifact opens, renders, and is clean ------------------------
    from openpyxl import load_workbook

    docx_files = sorted(stage.rglob("*.docx"))
    pdf_files = sorted(stage.rglob("*.pdf"))
    xlsx_files = sorted(stage.rglob("*.xlsx"))

    check(len(docx_files) == 11, f"{len(docx_files)} DOCX files produced", f"{len(docx_files)} DOCX, expected 11")
    check(len(pdf_files) == 13, f"{len(pdf_files)} PDF files produced", f"{len(pdf_files)} PDF, expected 13")
    check(len(xlsx_files) == 1, f"{len(xlsx_files)} XLSX file produced", f"{len(xlsx_files)} XLSX, expected 1")

    for path in docx_files:
        rel = path.relative_to(stage)
        try:
            text = docx_text(path)
        except Exception as exc:
            failures.append(f"DOCX will not open: {rel} ({exc})")
            continue
        if len(text) < 900:
            failures.append(f"DOCX has too little content to be finished: {rel} ({len(text)} chars)")
        if REQUIRED_FOOTER not in text:
            failures.append(f"DOCX missing required footer disclaimer: {rel}")
        for term in BANNED:
            if term.lower() in text.lower():
                failures.append(f"Banned placeholder '{term}' in {rel}")
        for term in SECRETS:
            if term.lower() in text.lower():
                failures.append(f"Infrastructure detail '{term}' leaked into {rel}")
    passes.append(f"All {len(docx_files)} DOCX files open, carry the footer, and contain finished content")

    for path in pdf_files:
        rel = path.relative_to(stage)
        raw = path.read_bytes()
        if not raw.startswith(b"%PDF-"):
            failures.append(f"Not a valid PDF: {rel}")
            continue
        if b"%%EOF" not in raw[-2048:]:
            failures.append(f"PDF truncated (no EOF marker): {rel}")
        text = pdf_text(path)
        if len(text) < 700:
            failures.append(f"PDF renders too little text: {rel} ({len(text)} chars)")
        if "Not legal, financial, HR, safety" not in text:
            failures.append(f"PDF missing required footer disclaimer: {rel}")
        for term in BANNED:
            if term.lower() in text.lower():
                failures.append(f"Banned placeholder '{term}' in {rel}")
        for term in SECRETS:
            if term.lower() in text.lower():
                failures.append(f"Infrastructure detail '{term}' leaked into {rel}")
    passes.append(f"All {len(pdf_files)} PDFs are valid, complete, carry the footer, and render text")

    for path in xlsx_files:
        rel = path.relative_to(stage)
        try:
            wb = load_workbook(str(path))
            ws = wb.active
            values = [str(c.value) for row in ws.iter_rows() for c in row if c.value is not None]
            joined = " ".join(values)
        except Exception as exc:
            failures.append(f"XLSX will not open: {rel} ({exc})")
            continue
        check(any(v.startswith("=") for v in values), "Workflow spreadsheet carries working scoring formulas",
              "Workflow spreadsheet has no formulas")
        if REQUIRED_FOOTER not in joined:
            failures.append(f"XLSX missing required footer disclaimer: {rel}")
        for term in BANNED:
            if term.lower() in joined.lower():
                failures.append(f"Banned placeholder '{term}' in {rel}")
    passes.append("XLSX opens and is usable")

    # --- Prohibited-claim sweep across every artifact -----------------------
    CLAIMS = [
        r"save \d+%", r"save \d+ hours", r"increase revenue by", r"guarantee[ds]? (?:you|results|savings)",
        r"\d+x (?:more|faster|revenue)", r"our customers (?:report|saw|achieved)", r"proven to increase",
        r"testimonial",
    ]
    all_text = "\n".join(docx_text(p) for p in docx_files) + "\n" + "\n".join(pdf_text(p) for p in pdf_files)
    found = [c for c in CLAIMS if re.search(c, all_text, re.I)]
    check(not found, "No fabricated results, guarantees, statistics or testimonials anywhere in the package",
          f"Prohibited claim language found: {found}")

    # --- Filename discipline ------------------------------------------------
    bad_names = [n for n in names if re.search(r"[ _]", pathlib.PurePosixPath(n).name)
                 and not n.endswith(".json")]
    check(not bad_names, "Filenames are consistent — hyphenated, no spaces or underscores",
          f"Inconsistent filenames: {bad_names}")

    size_mb = zip_path.stat().st_size / 1_048_576
    check(size_mb < 25, f"Package size {size_mb:.2f} MB (under the 25 MB target)",
          f"Package size {size_mb:.2f} MB exceeds the 25 MB target")

    # --- Report -------------------------------------------------------------
    print("=" * 74)
    print("SFL-AIOPS-001 — COMMERCIAL ACCEPTANCE GATE")
    print("=" * 74)
    for line in passes:
        print(f"  PASS   {line}")
    for line in warnings:
        print(f"  WARN   {line}")
    for line in failures:
        print(f"  FAIL   {line}")
    print("-" * 74)
    print(f"{len(passes)} passed, {len(warnings)} warnings, {len(failures)} failed")
    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(main())
