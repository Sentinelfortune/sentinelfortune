#!/usr/bin/env python3
"""
Build the SFL-AIOPS-001 downloadable product from its structured source.

SOURCE OF TRUTH
    product-source/SFL-AIOPS-001/product.json      product metadata
    product-source/SFL-AIOPS-001/content/*.json    one file per document
    product-source/SFL-AIOPS-001/content/prompts.json   the prompt library data

Content is stored as format-neutral block JSON, never as DOCX or PDF. Every
output format is rendered from the same blocks, so there is exactly one place
to edit a sentence. That is also what makes the product importable by an
automated production pipeline later: the pipeline consumes the JSON and reuses
these renderers rather than re-authoring the product.

BLOCK SCHEMA
    {"t":"h1"|"h2"|"h3", "x": str}      headings
    {"t":"p",  "x": str}                paragraph
    {"t":"ul"|"ol", "x": [str, ...]}    lists
    {"t":"table", "head":[str], "rows":[[str]], "widths":[float]}
    {"t":"note"|"warn", "x": str}       callout boxes
    {"t":"code", "x": str}              copy-ready prompt text (monospace)
    {"t":"kv", "rows": [[str, str]]}    two-column definition table
    {"t":"fields", "x": [str, ...]}     printed fill-in lines
    {"t":"hr"}                          rule
    {"t":"pb"}                          page break

Usage:  python3 scripts/build_product_aiops.py [--outdir DIR]
"""

import argparse
import json
import pathlib
import shutil
import sys
import zipfile
from datetime import datetime, timezone

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    HRFlowable,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

REPO = pathlib.Path(__file__).resolve().parent.parent
SRC = REPO / "product-source" / "SFL-AIOPS-001"

NAVY = colors.HexColor("#0e1a33")
GOLD = colors.HexColor("#9a7f2e")
GREY = colors.HexColor("#4a5468")
LIGHT = colors.HexColor("#f1eee8")
BORDER = colors.HexColor("#cfc9be")
WARN_BG = colors.HexColor("#fbf1dc")
WARN_LINE = colors.HexColor("#8f6310")
NOTE_BG = colors.HexColor("#e7f0f8")
NOTE_LINE = colors.HexColor("#1f5b8f")

DOCX_NAVY = RGBColor(0x0E, 0x1A, 0x33)
DOCX_GOLD = RGBColor(0x7D, 0x66, 0x1F)
DOCX_GREY = RGBColor(0x4A, 0x54, 0x68)

FOOTER_LINES = [
    "AI Operations Playbook & Toolkit",
    "Version 1.0",
    "Not legal, financial, HR, safety, code, permit, or regulatory advice.",
]


# ---------------------------------------------------------------------------
# Prompt library expansion
# ---------------------------------------------------------------------------

CATEGORY_ORDER = [
    ("field-notes", "Section 1 — Field-Note Cleanup"),
    ("estimates", "Section 2 — Estimate & Scope Drafting"),
    ("customer-comms", "Section 3 — Customer Communication"),
    ("sop", "Section 4 — SOP Extraction"),
    ("triage", "Section 5 — Voicemail & Email Triage"),
    ("csr", "Section 6 — CSR Objection Handling"),
]


def prompt_blocks(prompts):
    """Expand the prompt data file into document blocks."""
    blocks = []
    for idx, (key, heading) in enumerate(CATEGORY_ORDER):
        group = [p for p in prompts if p["category"] == key]
        if not group:
            continue
        if idx > 0:
            blocks.append({"t": "pb"})
        blocks.append({"t": "h2", "x": heading})
        blocks.append({"t": "p", "x": CATEGORY_INTRO[key]})
        for p in group:
            blocks.append({"t": "h3", "x": f"{p['id']} — {p['title']}"})
            blocks.append({"t": "kv", "rows": [
                ["Purpose", p["purpose"]],
                ["When to use", p["when"]],
                ["You must supply", "; ".join(p["inputs"])],
            ]})
            blocks.append({"t": "p", "x": "Copy-ready prompt:"})
            blocks.append({"t": "code", "x": p["prompt"]})
            if p.get("example"):
                blocks.append({"t": "p", "x": f"Example: {p['example']}"})
            blocks.append({"t": "warn", "x": p["warning"]})
    return blocks


CATEGORY_INTRO = {
    "field-notes": (
        "Field notes are where invoicing disputes begin. A technician writing on a phone in a crawlspace "
        "produces fragments; the office needs a record that supports an invoice and survives a callback six "
        "months later. These prompts reconstruct a usable record from fragments without inventing anything — "
        "every one of them requires the AI to mark what it does not know rather than fill the gap."
    ),
    "estimates": (
        "Estimates lose money in what they leave out. These prompts enforce a consistent structure: what is "
        "included, what is explicitly excluded, what conditions could change the scope, and what is routed to "
        "a licensed decision. None of them price anything — you insert your own figures after the draft exists."
    ),
    "customer-comms": (
        "One company should sound like one company regardless of who is typing. These prompts produce customer "
        "messages in a consistent voice, with your warranty and payment terms carried through as placeholders "
        "you fill from your Business Customization Worksheet."
    ),
    "sop": (
        "The most valuable procedures in a home-service business are undocumented and live in the owner's head. "
        "These prompts extract them by interview — the AI asks, you answer in plain speech, and the output is a "
        "structured procedure someone else can follow."
    ),
    "triage": (
        "The daily backlog of voicemails and emails is a sorting problem, not a writing problem. These prompts "
        "convert a pile of messages into a ranked action list with the urgent-safety items surfaced first and "
        "routed to a human immediately."
    ),
    "csr": (
        "Whoever answers the phone is negotiating on your behalf. These prompts prepare consistent, honest "
        "responses to the objections that actually come up, without discounting, without disparaging "
        "competitors, and without quoting a price the CSR is not authorised to give."
    ),
}


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------

def _shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _docx_footer(doc):
    for section in doc.sections:
        footer = section.footer
        # python-docx gives a footer with one empty paragraph; reuse it.
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(" · ".join(FOOTER_LINES))
        run.font.size = Pt(7.5)
        run.font.color.rgb = DOCX_GREY


def render_docx(doc_spec, blocks, out_path):
    doc = Document()

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)

    # Title block
    title = doc.add_paragraph()
    run = title.add_run(doc_spec["title"])
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = DOCX_NAVY
    title.paragraph_format.space_after = Pt(2)

    if doc_spec.get("subtitle"):
        sub = doc.add_paragraph()
        srun = sub.add_run(doc_spec["subtitle"])
        srun.font.size = Pt(11)
        srun.font.color.rgb = DOCX_GREY
        sub.paragraph_format.space_after = Pt(14)

    for b in blocks:
        t = b["t"]
        if t in ("h1", "h2", "h3"):
            size = {"h1": 16, "h2": 13.5, "h3": 11.5}[t]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14 if t != "h3" else 10)
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run(b["x"])
            r.font.size = Pt(size)
            r.font.bold = True
            r.font.color.rgb = DOCX_NAVY if t != "h3" else DOCX_GOLD
        elif t == "p":
            doc.add_paragraph(b["x"])
        elif t in ("ul", "ol"):
            style = "List Bullet" if t == "ul" else "List Number"
            for item in b["x"]:
                doc.add_paragraph(item, style=style)
        elif t == "kv":
            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for k, v in b["rows"]:
                cells = table.add_row().cells
                cells[0].width = Inches(1.5)
                cells[1].width = Inches(4.9)
                kr = cells[0].paragraphs[0].add_run(k)
                kr.font.bold = True
                kr.font.size = Pt(9.5)
                _shade(cells[0], "F1EEE8")
                vr = cells[1].paragraphs[0].add_run(v)
                vr.font.size = Pt(9.5)
            doc.add_paragraph()
        elif t == "table":
            table = doc.add_table(rows=1, cols=len(b["head"]))
            table.style = "Table Grid"
            for i, h in enumerate(b["head"]):
                cell = table.rows[0].cells[i]
                r = cell.paragraphs[0].add_run(h)
                r.font.bold = True
                r.font.size = Pt(9)
                _shade(cell, "0E1A33")
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for row in b["rows"]:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    r = cells[i].paragraphs[0].add_run(val)
                    r.font.size = Pt(9)
            doc.add_paragraph()
        elif t in ("note", "warn"):
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.rows[0].cells[0]
            _shade(cell, "FBF1DC" if t == "warn" else "E7F0F8")
            label = "REVIEW REQUIRED — " if t == "warn" else "NOTE — "
            p = cell.paragraphs[0]
            lr = p.add_run(label)
            lr.font.bold = True
            lr.font.size = Pt(9)
            lr.font.color.rgb = RGBColor(0x8F, 0x63, 0x10) if t == "warn" else RGBColor(0x1F, 0x5B, 0x8F)
            tr = p.add_run(b["x"])
            tr.font.size = Pt(9)
            doc.add_paragraph()
        elif t == "code":
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.rows[0].cells[0]
            _shade(cell, "F7F5F1")
            first = True
            for line in b["x"].split("\n"):
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                first = False
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(line)
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            doc.add_paragraph()
        elif t == "fields":
            for label in b["x"]:
                p = doc.add_paragraph()
                r = p.add_run(f"{label}: ")
                r.font.bold = True
                r.font.size = Pt(10)
                p.add_run("_" * max(12, 78 - len(label)))
        elif t == "hr":
            p = doc.add_paragraph()
            pr = p._p.get_or_add_pPr()
            bd = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "CFC9BE")
            bd.append(bottom)
            pr.append(bd)
        elif t == "pb":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    _docx_footer(doc)
    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# PDF rendering
# ---------------------------------------------------------------------------

def _pdf_styles():
    ss = getSampleStyleSheet()
    base = dict(fontName="Helvetica", fontSize=9.8, leading=14.2, textColor=colors.HexColor("#1a2338"))
    return {
        "title": ParagraphStyle("t", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=21,
                                leading=25, textColor=NAVY, spaceAfter=3),
        "subtitle": ParagraphStyle("st", parent=ss["Normal"], fontName="Helvetica", fontSize=10.5,
                                   leading=14, textColor=GREY, spaceAfter=16),
        "h1": ParagraphStyle("h1", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=15,
                             leading=19, textColor=NAVY, spaceBefore=15, spaceAfter=6),
        "h2": ParagraphStyle("h2", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=12.5,
                             leading=16, textColor=NAVY, spaceBefore=13, spaceAfter=5),
        "h3": ParagraphStyle("h3", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=10.8,
                             leading=14, textColor=GOLD, spaceBefore=11, spaceAfter=4),
        "p": ParagraphStyle("p", parent=ss["Normal"], spaceAfter=7, alignment=TA_LEFT, **base),
        "li": ParagraphStyle("li", parent=ss["Normal"], spaceAfter=3, **base),
        "cell": ParagraphStyle("c", parent=ss["Normal"], fontName="Helvetica", fontSize=8.6,
                               leading=11.6, textColor=colors.HexColor("#1a2338")),
        "cellb": ParagraphStyle("cb", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=8.6,
                                leading=11.6, textColor=NAVY),
        "cellh": ParagraphStyle("ch", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=8.6,
                                leading=11.6, textColor=colors.white),
        "code": ParagraphStyle("code", parent=ss["Normal"], fontName="Courier", fontSize=8.3,
                               leading=11.4, textColor=colors.HexColor("#12203a")),
        "callout": ParagraphStyle("co", parent=ss["Normal"], fontName="Helvetica", fontSize=8.8,
                                  leading=12.2, textColor=colors.HexColor("#1a2338")),
    }


def _footer_painter(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(BORDER)
    canvas.setLineWidth(0.5)
    canvas.line(0.85 * inch, 0.72 * inch, LETTER[0] - 0.85 * inch, 0.72 * inch)
    canvas.setFont("Helvetica", 6.8)
    canvas.setFillColor(GREY)
    canvas.drawString(0.85 * inch, 0.58 * inch, " · ".join(FOOTER_LINES))
    canvas.drawRightString(LETTER[0] - 0.85 * inch, 0.58 * inch, str(canvas.getPageNumber()))
    canvas.restoreState()


def _esc(text):
    return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def render_pdf(doc_spec, blocks, out_path):
    S = _pdf_styles()
    frame = Frame(0.85 * inch, 0.95 * inch, LETTER[0] - 1.7 * inch, LETTER[1] - 1.75 * inch, id="body")
    template = PageTemplate(id="main", frames=[frame], onPage=_footer_painter)
    pdf = BaseDocTemplate(str(out_path), pagesize=LETTER, pageTemplates=[template],
                          title=doc_spec["title"], author="Sentinel Fortune LLC",
                          subject="AI Operations Playbook & Toolkit v1.0")

    story = [Paragraph(_esc(doc_spec["title"]), S["title"])]
    if doc_spec.get("subtitle"):
        story.append(Paragraph(_esc(doc_spec["subtitle"]), S["subtitle"]))
    story.append(HRFlowable(width="100%", thickness=1.4, color=GOLD, spaceAfter=12))

    content_w = LETTER[0] - 1.7 * inch

    for b in blocks:
        t = b["t"]
        if t in ("h1", "h2", "h3"):
            story.append(Paragraph(_esc(b["x"]), S[t]))
        elif t == "p":
            story.append(Paragraph(_esc(b["x"]), S["p"]))
        elif t in ("ul", "ol"):
            items = [ListItem(Paragraph(_esc(i), S["li"]), leftIndent=14) for i in b["x"]]
            story.append(ListFlowable(items, bulletType="bullet" if t == "ul" else "1",
                                      leftIndent=16, bulletFontSize=7, spaceAfter=8))
        elif t == "kv":
            rows = [[Paragraph(_esc(k), S["cellb"]), Paragraph(_esc(v), S["cell"])] for k, v in b["rows"]]
            tbl = Table(rows, colWidths=[1.35 * inch, content_w - 1.35 * inch])
            tbl.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, BORDER),
                ("BACKGROUND", (0, 0), (0, -1), LIGHT),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.extend([tbl, Spacer(1, 9)])
        elif t == "table":
            head = [Paragraph(_esc(h), S["cellh"]) for h in b["head"]]
            body = [[Paragraph(_esc(c), S["cell"]) for c in r] for r in b["rows"]]
            widths = b.get("widths")
            col_w = [content_w * w for w in widths] if widths else [content_w / len(b["head"])] * len(b["head"])
            tbl = Table([head] + body, colWidths=col_w, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, BORDER),
                ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#faf9f6")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.extend([tbl, Spacer(1, 9)])
        elif t in ("note", "warn"):
            label = "REVIEW REQUIRED — " if t == "warn" else "NOTE — "
            colour = WARN_LINE if t == "warn" else NOTE_LINE
            bg = WARN_BG if t == "warn" else NOTE_BG
            para = Paragraph(f'<font color="{colour.hexval()}"><b>{label}</b></font>{_esc(b["x"])}', S["callout"])
            tbl = Table([[para]], colWidths=[content_w])
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), bg),
                ("BOX", (0, 0), (-1, -1), 0.5, colour),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.extend([tbl, Spacer(1, 9)])
        elif t == "code":
            lines = [Paragraph(_esc(l) if l.strip() else "&nbsp;", S["code"]) for l in b["x"].split("\n")]
            tbl = Table([[lines]], colWidths=[content_w])
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f7f5f1")),
                ("BOX", (0, 0), (-1, -1), 0.5, BORDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]))
            story.extend([tbl, Spacer(1, 9)])
        elif t == "fields":
            rows = [[Paragraph(f"<b>{_esc(l)}</b>", S["cell"]), Paragraph("", S["cell"])] for l in b["x"]]
            tbl = Table(rows, colWidths=[2.1 * inch, content_w - 2.1 * inch], rowHeights=[0.32 * inch] * len(rows))
            tbl.setStyle(TableStyle([
                ("LINEBELOW", (1, 0), (1, -1), 0.5, BORDER),
                ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.extend([tbl, Spacer(1, 9)])
        elif t == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceBefore=6, spaceAfter=10))
        elif t == "pb":
            story.append(PageBreak())

    pdf.build(story)


# ---------------------------------------------------------------------------
# XLSX (workflow scoring worksheet)
# ---------------------------------------------------------------------------

def render_workflow_xlsx(out_path):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Workflow Scoring"

    navy = PatternFill("solid", fgColor="0E1A33")
    light = PatternFill("solid", fgColor="F1EEE8")
    thin = Side(style="thin", color="CFC9BE")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws["A1"] = "Operations Workflow Scoring Worksheet"
    ws["A1"].font = Font(size=15, bold=True, color="0E1A33")
    ws["A2"] = ("Score each workflow 1-5. Priority is calculated automatically: "
                "(Frequency x Time Cost) - (Risk x 2). Work the highest scores first.")
    ws["A2"].font = Font(size=9, color="4A5468")
    ws["A2"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells("A2:G2")
    ws.row_dimensions[2].height = 28

    headers = ["Workflow", "Who owns it today", "Frequency (1-5)", "Time cost (1-5)",
               "Risk if AI-assisted (1-5)", "Priority score", "Notes"]
    for i, h in enumerate(headers, start=1):
        c = ws.cell(row=4, column=i, value=h)
        c.font = Font(bold=True, color="FFFFFF", size=9)
        c.fill = navy
        c.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
        c.border = border
    ws.row_dimensions[4].height = 34

    workflows = [
        "Inbound call handling and booking",
        "Dispatch and scheduling",
        "On-site diagnosis write-up",
        "Field notes to office handoff",
        "Estimate and scope drafting",
        "Invoice preparation",
        "Post-job customer follow-up",
        "Callback and complaint handling",
        "New technician onboarding",
        "Maintenance-plan renewals",
    ]
    for r, name in enumerate(workflows, start=5):
        ws.cell(row=r, column=1, value=name).border = border
        ws.cell(row=r, column=1).alignment = Alignment(wrap_text=True, vertical="center")
        for col in range(2, 6):
            cell = ws.cell(row=r, column=col)
            cell.border = border
            cell.fill = light
            cell.alignment = Alignment(horizontal="center", vertical="center")
        pc = ws.cell(row=r, column=6, value=f"=IF(COUNT(C{r}:E{r})=3,(C{r}*D{r})-(E{r}*2),\"\")")
        pc.border = border
        pc.font = Font(bold=True, color="0E1A33")
        pc.alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row=r, column=7).border = border
        ws.row_dimensions[r].height = 26

    guide_row = len(workflows) + 7
    ws.cell(row=guide_row, column=1, value="Scoring guide").font = Font(bold=True, size=11, color="0E1A33")
    for offset, text in enumerate([
        "Frequency — 1 = monthly or less, 3 = weekly, 5 = many times every day.",
        "Time cost — 1 = a couple of minutes, 3 = 15-30 minutes, 5 = an hour or more each time.",
        "Risk if AI-assisted — 1 = internal only, 3 = customer-facing, 5 = touches pricing, code, permits or safety.",
        "Risk is weighted double on purpose. A high-risk workflow needs a human decision, not a faster draft.",
        "Anything scoring 5 on risk stays with a qualified person. Do not automate it because the score looks good.",
    ], start=1):
        c = ws.cell(row=guide_row + offset, column=1, value=text)
        c.font = Font(size=9, color="4A5468")
        ws.merge_cells(start_row=guide_row + offset, start_column=1, end_row=guide_row + offset, end_column=7)

    footer_row = guide_row + 7
    fc = ws.cell(row=footer_row, column=1, value=" · ".join(FOOTER_LINES))
    fc.font = Font(size=7.5, color="6F7889")
    ws.merge_cells(start_row=footer_row, start_column=1, end_row=footer_row, end_column=7)

    for col, width in zip("ABCDEFG", [34, 20, 12, 12, 14, 12, 30]):
        ws.column_dimensions[col].width = width
    ws.freeze_panes = "A5"

    wb.save(str(out_path))


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def load_json(path):
    with open(path, encoding="utf-8") as fh:
        return json.load(fh)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--outdir", default=str(REPO / "dist" / "product"))
    args = parser.parse_args()

    product = load_json(SRC / "product.json")
    prompts = load_json(SRC / "content" / "prompts.json")

    out_root = pathlib.Path(args.outdir)
    stage = out_root / product["package_stem"]
    if stage.exists():
        shutil.rmtree(stage)
    stage.mkdir(parents=True)

    manifest_files = []

    for spec in product["documents"]:
        blocks = load_json(SRC / "content" / spec["source"])["blocks"]
        if spec.get("inject") == "prompts":
            marker = next(i for i, b in enumerate(blocks) if b.get("t") == "inject")
            blocks = blocks[:marker] + prompt_blocks(prompts) + blocks[marker + 1:]

        target_dir = stage / spec["folder"] if spec["folder"] else stage
        target_dir.mkdir(parents=True, exist_ok=True)

        for fmt in spec["formats"]:
            path = target_dir / f"{spec['filename']}.{fmt}"
            if fmt == "pdf":
                render_pdf(spec, blocks, path)
            elif fmt == "docx":
                render_docx(spec, blocks, path)
            rel = str(path.relative_to(stage))
            manifest_files.append({
                "asset_id": spec["id"],
                "title": spec["title"],
                "path": rel,
                "format": fmt,
                "bytes": path.stat().st_size,
            })
            print(f"  {rel}  ({path.stat().st_size:,} bytes)")

    # Workflow scoring spreadsheet
    xlsx_spec = product["spreadsheet"]
    xlsx_dir = stage / xlsx_spec["folder"]
    xlsx_dir.mkdir(parents=True, exist_ok=True)
    xlsx_path = xlsx_dir / f"{xlsx_spec['filename']}.xlsx"
    render_workflow_xlsx(xlsx_path)
    manifest_files.append({
        "asset_id": xlsx_spec["id"],
        "title": xlsx_spec["title"],
        "path": str(xlsx_path.relative_to(stage)),
        "format": "xlsx",
        "bytes": xlsx_path.stat().st_size,
    })
    print(f"  {xlsx_path.relative_to(stage)}  ({xlsx_path.stat().st_size:,} bytes)")

    # PRODUCT-MANIFEST.json is the governed import contract consumed by the Shop
    # Admin (shop-worker/src/lib/product-manifest.ts). Field names are camelCase
    # because that is the contract; see PRODUCT_MANIFEST_SCHEMA.md.
    listing = product["listing"]
    # Cover image — copied into the package root and declared, so the governed
    # import can attach it as the product's COVER without a separate upload.
    cover_spec = product.get("cover_image")
    if cover_spec:
        cover_src = SRC / cover_spec["source"]
        if not cover_src.exists():
            raise SystemExit(f"Cover image missing: {cover_src}. Run scripts/build_cover_aiops.py first.")
        cover_dest = stage / cover_spec["package_path"]
        cover_dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(cover_src, cover_dest)
        manifest_files.append({
            "asset_id": "COVER",
            "title": cover_spec.get("title", "Product cover"),
            "path": cover_spec["package_path"],
            "format": cover_spec.get("format", "png"),
            "bytes": cover_dest.stat().st_size,
        })
        print(f"  {cover_spec['package_path']}  ({cover_dest.stat().st_size:,} bytes)")

    manifest = {
        "contractVersion": product["contract_version"],
        "producer": product["producer"],
        "builtAt": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),

        "sku": product["sku"],
        "slug": product["slug"],
        "title": product["title"],
        "version": product["version"],
        "edition": product["edition"],
        "category": product["category"],
        "audience": product["audience"],
        "licenseType": product["license_type"],
        "licenseName": product["license"],
        "supportedFormats": product["supported_formats"],

        "shortDescription": listing["short_description"],
        "problemSolved": listing["problem_solved"],
        "description": listing["description"],
        "deliverables": listing["deliverables"],
        "notIncluded": listing["not_included"],
        "faqs": listing["faqs"],
        "responsibleUseText": listing["responsible_use_text"],
        "refundEligible": product["refund_eligible"],
        "refundPolicySummary": listing["refund_policy_summary"],

        "recommendedPriceCents": product["recommended_price_cents"],
        "currency": product["currency"],
        "downloadLinkExpiryHours": product["download_link_expiry_hours"],
        "maxDownloads": product["max_downloads"],
        "coverImage": (product.get("cover_image") or {}).get("package_path"),

        # Production metadata — informational, not part of the import contract.
        "promptCount": len(prompts),
        "promptCategories": {key: sum(1 for p in prompts if p["category"] == key)
                             for key, _ in CATEGORY_ORDER},
        "assetCount": len(product["documents"]) + 1,
        "fileCount": len(manifest_files),

        "files": sorted(manifest_files, key=lambda f: f["path"]),
    }
    with open(stage / "PRODUCT-MANIFEST.json", "w", encoding="utf-8") as fh:
        json.dump(manifest, fh, indent=2)
        fh.write("\n")

    zip_path = out_root / f"{product['package_stem']}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(stage.rglob("*")):
            if path.is_file():
                zf.write(path, str(path.relative_to(stage)))

    print(f"\nPackage: {zip_path}  ({zip_path.stat().st_size:,} bytes)")
    print(f"Assets: {manifest['assetCount']}   Files: {manifest['fileCount']}   "
          f"Prompts: {manifest['promptCount']}   Contract: v{manifest['contractVersion']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
