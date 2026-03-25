import re
import logging
from io import BytesIO
from xml.sax.saxutils import escape as xml_escape

logger = logging.getLogger(__name__)

_LABEL_RE = re.compile(r"^([A-Z][A-Za-z0-9\s/,\-]{0,45}):\s*(.*)")
_BULLET_RE = re.compile(r"^[•\-]\s+")
_NUMBERED_RE = re.compile(r"^\d+[\.\)]\s+")


def _detect_label(line: str):
    m = _LABEL_RE.match(line)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return None, None


def clean_text(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    return text


def _safe_filename(title: str, ext: str) -> str:
    clean = re.sub(r"[^\w\s\-]", "", title)
    clean = re.sub(r"\s+", "_", clean.strip())[:45]
    return f"{clean or 'document'}.{ext}"


def generate_pdf(title: str, content_type: str, body: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    body = clean_text(body)
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=title,
        author="Sentinel Fortune",
    )

    sf_dark = colors.HexColor("#3A3A3A")
    sf_mid = colors.HexColor("#5A5A5A")
    sf_grey = colors.HexColor("#909090")

    title_style = ParagraphStyle(
        "SFTitle",
        fontName="Helvetica-Bold",
        fontSize=20,
        spaceAfter=4,
        textColor=sf_dark,
        alignment=TA_LEFT,
    )
    subtitle_style = ParagraphStyle(
        "SFSubtitle",
        fontName="Helvetica",
        fontSize=10,
        spaceAfter=14,
        textColor=sf_grey,
    )
    label_style = ParagraphStyle(
        "SFLabel",
        fontName="Helvetica-Bold",
        fontSize=11,
        spaceBefore=12,
        spaceAfter=3,
        textColor=sf_mid,
    )
    label_inline_style = ParagraphStyle(
        "SFLabelInline",
        fontName="Helvetica",
        fontSize=11,
        spaceAfter=6,
        leading=17,
    )
    body_style = ParagraphStyle(
        "SFBody",
        fontName="Helvetica",
        fontSize=11,
        spaceAfter=6,
        leading=17,
    )
    bullet_style = ParagraphStyle(
        "SFBullet",
        fontName="Helvetica",
        fontSize=11,
        spaceAfter=4,
        leading=17,
        leftIndent=18,
    )

    story = []
    story.append(Paragraph(xml_escape(title), title_style))
    story.append(Paragraph(f"Sentinel Fortune — {xml_escape(content_type)}", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=sf_grey, spaceAfter=14))

    blank_count = 0
    for line in body.split("\n"):
        line = line.strip()

        if not line:
            blank_count += 1
            if blank_count <= 1:
                story.append(Spacer(1, 0.18 * cm))
            continue
        blank_count = 0

        label, content = _detect_label(line)
        if label:
            if content:
                para = Paragraph(
                    f"<b>{xml_escape(label)}:</b> {xml_escape(content)}",
                    label_inline_style,
                )
                story.append(para)
            else:
                story.append(Paragraph(xml_escape(label) + ":", label_style))
        elif _BULLET_RE.match(line):
            clean = _BULLET_RE.sub("", line)
            story.append(Paragraph(f"• {xml_escape(clean)}", bullet_style))
        elif _NUMBERED_RE.match(line):
            story.append(Paragraph(xml_escape(line), bullet_style))
        else:
            story.append(Paragraph(xml_escape(line), body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx(title: str, content_type: str, body: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm

    body = clean_text(body)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(58, 58, 58)

    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run(f"Sentinel Fortune — {content_type}")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(144, 144, 144)

    doc.add_paragraph()

    for line in body.split("\n"):
        line = line.strip()
        if not line:
            continue

        label, content = _detect_label(line)
        if label:
            if content:
                p = doc.add_paragraph()
                r_label = p.add_run(f"{label}: ")
                r_label.bold = True
                r_label.font.size = Pt(11)
                r_content = p.add_run(content)
                r_content.font.size = Pt(11)
            else:
                p = doc.add_paragraph()
                r = p.add_run(f"{label}:")
                r.bold = True
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(58, 58, 58)
        elif _BULLET_RE.match(line):
            clean = _BULLET_RE.sub("", line)
            try:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(clean).font.size = Pt(11)
            except Exception:
                p = doc.add_paragraph()
                p.add_run(f"• {clean}").font.size = Pt(11)
        elif _NUMBERED_RE.match(line):
            p = doc.add_paragraph()
            p.add_run(line).font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            p.add_run(line).font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def safe_filename(title: str, ext: str) -> str:
    return _safe_filename(title, ext)
